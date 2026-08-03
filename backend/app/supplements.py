"""补充资料预检、时点隔离、独立证据提取与本地留存。

补充资料必须绑定一个已经存在的父运行，不能脱离案例单独生效。
资料日期与原案例时点分别保存，补充日期不会改写原始时点。
任何文件都必须先确认授权和脱敏，否则只记录拒绝原因。
公开资料同样需要确认个人信息边界，不能因公开而跳过检查。
文件扩展名采用白名单，未知格式不会交给通用解析器猜测。
单个文件设置体积上限，防止补充入口成为任意文件存储通道。
文本内容扫描高风险个人信息，命中时要求脱敏后重新提交。
账龄、期后回款、信用政策和合同条款作为独立证据保存。
独立证据不会覆盖原年报中的营业收入或应收账款字段。
结构化字段更正仅为旧接口兼容，必须明确标为字段更正模式。
更正字段必须使用允许的字段编号，任意键值不会进入计算。
金额解析拒绝空值和非数字文本，不从自由描述中猜测金额。
结构化粘贴内容必须是有效 JSON，解析失败时不部分接受。
JSON 可以同时携带独立证据与旧字段更正，但两者分别存储。
账龄摘要保留原始结构，不替使用者推算客户级风险结论。
期后回款摘要保留比率或说明，不自动认定应收余额可收回。
信用政策变化只形成待验证背景，不能直接改变规则触发状态。
合同结算条款只作为解释证据，不替代收入准则的履约判断。
CSV 和表格只读取明确表头，不依赖列位置猜测字段含义。
无法识别为规则字段的结构化表仍可登记为待人工确认材料。
Word 只做基本可读性预检，不从自由文本自动抽取审计数字。
PDF 与文本文件可以登记来源，但金额仍需结构化字段支持。
每份补充资料生成独立编号、证据编号、文件哈希和保存时间。
原文件名经过安全清理，不能携带目录穿越或非法控制字符。
文件哈希用于确认版本，不说明文件来源真实或内容正确。
资料类型、说明和来源文件名均限制长度，避免日志被滥用。
补充证据的支持状态固定为待人工确认，不由解析成功自动提升。
解析成功只表示程序可读，不等于专业含义或口径已经签字。
补充资料必须绑定当前已接入规则，不能声明不存在的规则编号。
续分析只重跑绑定规则，避免无关规则消耗模型预算。
续分析复制父运行原字段，再显式加入补充证据或更正字段。
更正字段使用新的证据编号，原字段仍保留在父运行版本中。
父子运行通过补充资料编号和续分析模式建立版本链。
续分析默认重新执行 RAG 与三 Agent，不固定关闭模型检查。
如果案例禁止模型传输，续分析仍只能完成本地计算预检。
补充资料失败不会修改父运行，也不会删除此前成功的证据。
存储目录按运行命名空间隔离，测试资料不能污染演示记录。
读取补充资料时验证编号格式，避免把编号解释为文件路径。
未知编号返回未找到，不遍历或回显本机目录结构。
资料状态区分可续分析、仅登记和已拒绝，不能用单一成功标签。
仅有说明文字时可以登记材料，但必须保持待人工确认状态。
既无文件也无结构化内容时拒绝创建空补充记录。
高风险信息命中、解析失败和授权缺失会分别进入问题列表。
问题列表面向使用者说明修正方向，不包含原始敏感内容。
网页展示补充证据时必须同时显示资料日期和支持状态。
报告导出时补充证据只作为来源链的一部分，不冒充已执行程序。
新增资料类型时必须决定解析边界、证据编号和人工确认方式。
新增自动抽取时必须保留原文件、页码定位和人工复核入口。
本模块的核心原则是新增证据而不改写历史，并诚实保留缺口。
"""

from __future__ import annotations

import csv
import hashlib
import io
import json
import os
import re
import time
import uuid
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook

from .schemas import AI_GENERATED_CONTENT_NOTICE


ALLOWED_EXTENSIONS = {".pdf", ".xlsx", ".csv", ".json", ".docx", ".txt"}
ALLOWED_FIELDS = {
    "revenue_current",
    "revenue_previous",
    "ar_current",
    "ar_previous",
    "operating_cash_flow_current",
    "operating_cash_flow_previous",
    "net_profit_current",
}
MAX_FILE_BYTES = 15 * 1024 * 1024
HIGH_RISK_PATTERNS = {
    "居民身份证号": re.compile(r"(?<!\d)\d{17}[0-9Xx](?!\d)"),
    "中国大陆手机号": re.compile(r"(?<!\d)1[3-9]\d{9}(?!\d)"),
    "疑似银行卡号": re.compile(r"(?<!\d)\d{16,19}(?!\d)"),
}


def _directory(workspace_root: Path) -> Path:
    namespace = re.sub(r"[^A-Za-z0-9_-]", "", os.getenv("AUDITTRACE_RUNTIME_NAMESPACE", ""))
    base = workspace_root / "backend" / "runtime"
    path = (base / namespace if namespace else base) / "supplements"
    path.mkdir(parents=True, exist_ok=True)
    return path


def _safe_name(filename: str) -> str:
    name = Path(filename).name
    cleaned = re.sub(r"[^\w.\-（）()：： ]", "_", name, flags=re.UNICODE)
    return cleaned[:120] or "supplement.bin"


def _to_number(value: Any) -> float | None:
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        return float(value)
    if isinstance(value, str):
        clean = value.strip().replace(",", "").replace("，", "")
        if not clean:
            return None
        try:
            return float(clean)
        except ValueError:
            return None
    return None


def _normalize_mapping(data: Any) -> dict[str, float]:
    if not isinstance(data, dict):
        return {}
    result: dict[str, float] = {}
    for key, value in data.items():
        if key not in ALLOWED_FIELDS:
            continue
        number = _to_number(value)
        if number is not None:
            result[key] = number
    return result


def _normalize_supplement_evidence(data: Any) -> list[dict[str, Any]]:
    if not isinstance(data, dict):
        return []
    evidence: list[dict[str, Any]] = []
    mapping = {
        "aging_summary": "账龄结构",
        "subsequent_receipts_summary": "期后回款",
        "credit_policy_change": "信用政策",
        "contract_terms": "合同结算条款",
    }
    for key, label in mapping.items():
        value = data.get(key)
        if value in (None, "", [], {}):
            continue
        evidence.append(
            {
                "evidence_kind": key,
                "field_label": label,
                "details": value,
                "support_status": "pending_human_confirmation",
            }
        )
    return evidence


def _extract_structured(
    filename: str, content: bytes, structured_json: str | None
) -> tuple[dict[str, float], list[dict[str, Any]], list[str]]:
    issues: list[str] = []
    fields: dict[str, float] = {}
    supplemental_evidence: list[dict[str, Any]] = []
    if structured_json and structured_json.strip():
        try:
            parsed = json.loads(structured_json)
            # 兼容旧版顶层字段，也支持明确的 base_field_corrections 包装。
            fields.update(_normalize_mapping(parsed))
            if isinstance(parsed, dict):
                fields.update(_normalize_mapping(parsed.get("base_field_corrections")))
            supplemental_evidence.extend(_normalize_supplement_evidence(parsed))
        except json.JSONDecodeError:
            issues.append("结构化粘贴内容不是有效 JSON。")

    suffix = Path(filename).suffix.lower()
    try:
        if suffix == ".json" and content:
            parsed = json.loads(content.decode("utf-8-sig"))
            fields.update(_normalize_mapping(parsed))
            if isinstance(parsed, dict):
                fields.update(_normalize_mapping(parsed.get("base_field_corrections")))
            supplemental_evidence.extend(_normalize_supplement_evidence(parsed))
        elif suffix == ".csv":
            rows = list(csv.DictReader(io.StringIO(content.decode("utf-8-sig"))))
            for row in rows:
                field_id = (row.get("field_id") or row.get("字段") or "").strip()
                value = row.get("value") if "value" in row else row.get("金额")
                number = _to_number(value)
                if field_id in ALLOWED_FIELDS and number is not None:
                    fields[field_id] = number
            if rows and not fields:
                supplemental_evidence.append(
                    {
                        "evidence_kind": "structured_table",
                        "field_label": "补充资料结构化表",
                        "details": {"columns": list(rows[0]), "row_count": len(rows)},
                        "support_status": "pending_human_confirmation",
                    }
                )
        elif suffix == ".xlsx":
            workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            sheet = workbook.active
            rows = list(sheet.iter_rows(values_only=True))
            if rows:
                headers = [str(value or "").strip() for value in rows[0]]
                for values in rows[1:]:
                    row = dict(zip(headers, values))
                    field_id = str(row.get("field_id") or row.get("字段") or "").strip()
                    number = _to_number(row.get("value") if "value" in row else row.get("金额"))
                    if field_id in ALLOWED_FIELDS and number is not None:
                        fields[field_id] = number
                if len(rows) > 1 and not fields:
                    supplemental_evidence.append(
                        {
                            "evidence_kind": "structured_table",
                            "field_label": "补充资料结构化表",
                            "details": {"columns": headers, "row_count": len(rows) - 1},
                            "support_status": "pending_human_confirmation",
                        }
                    )
        elif suffix == ".docx":
            # Word 只做可读性预检；字段必须另行结构化粘贴，避免从自由文本猜金额。
            document = Document(io.BytesIO(content))
            if not any(paragraph.text.strip() for paragraph in document.paragraphs) and not document.tables:
                issues.append("Word 文件没有可读取的正文或表格。")
    except (UnicodeDecodeError, json.JSONDecodeError, OSError, ValueError, KeyError) as error:
        issues.append(f"文件解析失败：{type(error).__name__}")
    return fields, supplemental_evidence, issues


def create_supplement(
    workspace_root: Path,
    *,
    parent_run_id: str,
    material_type: str,
    authorized: bool,
    desensitized: bool,
    bound_rule_ids: list[str],
    as_of_date: str,
    note: str,
    filename: str,
    content: bytes,
    structured_json: str | None,
) -> dict[str, Any]:
    issues: list[str] = []
    if not authorized:
        issues.append("未确认资料已获授权。")
    if not desensitized:
        issues.append("未确认资料已经脱敏。")
    if not bound_rule_ids or any(rule not in {"R1", "R2"} for rule in bound_rule_ids):
        issues.append("必须绑定当前已接入的 R1 或 R2。")
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", as_of_date or ""):
        issues.append("资料日期必须使用 YYYY-MM-DD。")
    safe_name = _safe_name(filename or "structured.json")
    suffix = Path(safe_name).suffix.lower()
    if content and suffix not in ALLOWED_EXTENSIONS:
        issues.append("文件类型不支持；仅允许 PDF、XLSX、CSV、JSON、DOCX、TXT。")
    if len(content) > MAX_FILE_BYTES:
        issues.append("文件超过 15MB 限制。")
    if not content and not (structured_json or "").strip():
        issues.append("必须上传文件或粘贴结构化字段。")

    scan_text = ""
    if Path(safe_name).suffix.lower() in {".json", ".csv", ".txt"}:
        scan_text += content.decode("utf-8-sig", errors="ignore")
    scan_text += structured_json or ""
    for label, pattern in HIGH_RISK_PATTERNS.items():
        if pattern.search(scan_text):
            issues.append(f"检出{label}，请脱敏后重传。")

    structured_fields, supplemental_evidence, parse_issues = _extract_structured(safe_name, content, structured_json)
    issues.extend(parse_issues)
    supplement_id = f"SUP-{uuid.uuid4().hex[:12].upper()}"
    file_hash = hashlib.sha256(content).hexdigest().upper() if content else hashlib.sha256((structured_json or "").encode("utf-8")).hexdigest().upper()
    if not supplemental_evidence and not structured_fields and (content or note.strip()):
        supplemental_evidence.append(
            {
                "evidence_kind": "registered_material",
                "field_label": material_type.strip()[:100] or "其他补充资料",
                "details": {"note": note.strip()[:500], "filename": safe_name},
                "support_status": "pending_human_confirmation",
            }
        )
    for index, item in enumerate(supplemental_evidence, start=1):
        item["evidence_id"] = f"{supplement_id}-E{index:02d}"
        item["as_of_date"] = as_of_date
        item["source_file"] = safe_name
        item["file_sha256"] = file_hash
        item["source_mode"] = "supplement_independent"
    status = "ready_for_rerun" if not issues and (structured_fields or supplemental_evidence) else "registered_only" if not issues else "rejected"
    record = {
        "ai_generated_content_notice": AI_GENERATED_CONTENT_NOTICE,
        "supplement_id": supplement_id,
        "parent_run_id": parent_run_id,
        "material_type": material_type.strip()[:100],
        "authorized": authorized,
        "desensitized": desensitized,
        "bound_rule_ids": list(dict.fromkeys(bound_rule_ids)),
        "as_of_date": as_of_date,
        "note": note.strip()[:1000],
        "original_filename": safe_name,
        "file_sha256": file_hash,
        "file_size": len(content),
        "structured_fields": structured_fields,
        "structured_evidence": supplemental_evidence,
        "field_correction_mode": "legacy_explicit_only" if structured_fields else "none",
        "status": status,
        "issues": issues,
        "boundary": (
            "ready_for_rerun 只表示补充证据可被程序读取；原T0不会被资料日期改写，来源真实性、口径和专业含义仍须人工复核。"
            if status == "ready_for_rerun"
            else "文件已登记但没有可靠结构化字段，不会自动改写风险卡。"
        ),
        "created_at": time.strftime("%Y-%m-%dT%H:%M:%S%z"),
    }
    directory = _directory(workspace_root) / supplement_id
    directory.mkdir(parents=True, exist_ok=False)
    if content:
        (directory / safe_name).write_bytes(content)
    (directory / "record.json").write_text(json.dumps(record, ensure_ascii=False, indent=2), encoding="utf-8")
    return record


def load_supplement(workspace_root: Path, supplement_id: str) -> dict[str, Any] | None:
    if not re.fullmatch(r"SUP-[A-Z0-9]+", supplement_id):
        return None
    path = _directory(workspace_root) / supplement_id / "record.json"
    return json.loads(path.read_text(encoding="utf-8")) if path.exists() else None
