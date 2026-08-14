"""人工复核后的缓存、回放、版本链与报告导出限制。

报告是人工批准后的交付物，不是每次程序运行的自动副产品。
未找到运行记录时不能创建空报告或使用最近一次运行替代。
人工处理仍为未复核时，报告导出接口必须拒绝请求。
人工身份标为自动化时，不能获得正式导出批准资格。
复核人姓名为空时不能通过真人批准闸门。
导出批准复选框必须由真人明确提交，系统不能默认勾选。
人工处理与 AI 建议分别展示，报告不能把二者合并成结论。
程序筛查与 AI 草稿分别展示，读者能够识别各自来源。
报告标题明确属于审计计划阶段，不扩展为承接或续聘报告。
报告范围明确聚焦销售与收款循环，不冒充完整审计计划。
每项规则展示工程版本、计算指标、来源验证和程序状态。
R1 指标展示收入增速、应收增速、增速差和金额重要性边界。
净额或账面余额口径必须随来源字段写入报告。
三年趋势缺失时报告保留限制，不用空值推断持续期间。
计算字段列出证据编号、文档编号、页码和定位说明。
RAG 片段列出检索编号、页码、原文摘录和待回页状态。
补充证据列出资料日期、文件哈希和待人工确认状态。
三 Agent 步骤分别展示角色状态、失败阶段和稳定失败代码。
模型失败时不导出不存在的最终草稿，只展示诚实失败状态。
最终草稿存在时明确标记 AI 辅助生成和人工责任边界。
草稿中的主张、正常解释、资料缺口和待索取资料分栏呈现。
报告不得使用舞弊、造假或审计意见等越权定性表述。
人工处理记录包含处理状态、备注、复核人和复核时间。
人工复核时间来自提交记录，系统不能替真人生成签字日期。
报告文件名使用安全运行编号，不能包含外部输入路径字符。
导出目录只保存生成文档，不携带模型密钥、请求日志或缓存。
Word 样式保证中文字体可读，同时保留版本和证据编号的等宽视觉。
页眉页脚展示报告版本与 AI 边界，不使用过期工程版本。
表格过长时允许跨页，但不能裁掉证据编号或支持状态。
缓存只允许在真人批准后创建，未批准运行不能固化为交付缓存。
缓存内容保存公开运行结构，不保存环境变量和模型密钥。
缓存编号采用安全格式，读取时不能转换为任意文件路径。
缓存回放产生新的回放标记，但不重新执行模型或检索。
回放状态必须说明来自缓存，不能当作新的真实模型成功。
回放保留原始运行编号和版本链，支持追踪最初证据来源。
旧缓存缺少新字段时保持只读兼容，不伪造失败阶段或草稿。
报告版本升级时必须保留旧报告可读性和明确版本标识。
新增导出字段时应来自已存运行，不在导出阶段重新计算。
导出阶段不访问外部网络，避免报告内容随网络状态变化。
导出阶段不重新读取任意年报路径，只使用已登记证据元数据。
人工批准之后若来源文件变化，原运行哈希仍用于揭示版本差异。
报告生成成功不代表专业复核完成质量已得到系统评价。
报告是待核查事项备忘录，不是审计工作底稿的完整替代。
报告不得写成投资建议、信用评级或监管调查结论。
自动化测试生成的报告必须带自动化水印，不能对外冒充真人签发。
测试命名空间的缓存和报告与正式运行目录隔离。
删除缓存或报告不属于普通运行流程，当前模块不提供批量删除。
新增格式导出时必须复用人工批准、证据链和 AI 声明闸门。
回放必须引用原缓存编号和原运行编号，不能伪装成当前时点的新分析。
缓存读取再次校验安全编号，外部输入不能控制报告或缓存文件路径。
版本链只允许从已存在且已批准的报告派生，孤立子版本不会被创建。
新版本保存父版本编号和变更说明，读者能够还原交付演进关系。
版本变化不自动继承新的人工批准，修改内容后仍需明确复核边界。
报告构建读取结构化运行响应，不会重新查询模型补全缺失段落。
缺失证据字段在报告中显示限制，不能用空字符串制造完整表格假象。
文档中的 AI 提示必须与 API 固定声明一致，样式调整不能改变责任含义。
报告生成路径固定在受控导出目录，企业名称不能成为目录片段。
生成 Word 失败时不留下可下载的半成品，也不能登记一个成功报告版本。
页面预览和 Word 导出应引用同一运行结构，避免两个渠道出现不同结论。
正式采用仍需审计团队执行独立程序，本报告只服务计划阶段候选整理。
缓存、回放和导出都不能改变来源时点，未来材料不得进入旧运行报告。
人工批准证明允许交付当前版本，不代表其中每条 AI 主张已经事实认定。
交付模块的完整性目标是如实呈现证据和限制，而不是追求无空项的版面。
本模块的核心原则是先有人负责，再允许把运行结果固化交付。
"""

from __future__ import annotations

import json
import os
import re
import time
import uuid
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .schemas import AI_GENERATED_CONTENT_NOTICE, HumanReviewRequest, RunResponse, StoredRunResponse


REPORT_VERSION = "report_v2"


def _safe_id(value: str, prefix: str) -> bool:
    return bool(re.fullmatch(fr"{prefix}-[A-Z0-9-]+", value))


def _runtime_base(workspace_root: Path) -> Path:
    namespace = re.sub(r"[^A-Za-z0-9_-]", "", os.getenv("AUDITTRACE_RUNTIME_NAMESPACE", ""))
    base = workspace_root / "backend" / "runtime"
    return base / namespace if namespace else base


def _cache_dir(workspace_root: Path) -> Path:
    path = _runtime_base(workspace_root) / "cache"
    path.mkdir(parents=True, exist_ok=True)
    return path


def cache_run(workspace_root: Path, stored: StoredRunResponse) -> dict[str, Any]:
    review = stored.human_review
    if review is None or review.status == "未复核" or not review.export_approved:
        raise ValueError("只有完成人工复核并勾选允许导出的运行才能写入缓存。")
    if review.reviewer_type != "human":
        raise ValueError("自动化身份不能批准正式缓存。")
    cache_id = f"CACHE-{uuid.uuid4().hex[:12].upper()}"
    payload = {
        "ai_generated_content_notice": AI_GENERATED_CONTENT_NOTICE,
        "cache_id": cache_id,
        "cached_at": time.strftime("%Y-%m-%dT%H:%M:%S%z"),
        "source_run_id": stored.run.run_id,
        "source_schema_version": stored.run.schema_version,
        "boundary": "这是人工批准后的缓存回放源，不代表本次重新调用模型。",
        "stored": stored.model_dump(mode="json"),
    }
    (_cache_dir(workspace_root) / f"{cache_id}.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return {key: value for key, value in payload.items() if key != "stored"}


def load_cache(workspace_root: Path, cache_id: str) -> dict[str, Any] | None:
    if not _safe_id(cache_id, "CACHE"):
        return None
    path = _cache_dir(workspace_root) / f"{cache_id}.json"
    return json.loads(path.read_text(encoding="utf-8")) if path.exists() else None


def replay_cache(workspace_root: Path, cache_id: str) -> RunResponse | None:
    payload = load_cache(workspace_root, cache_id)
    if payload is None:
        return None
    run_data = deepcopy(payload["stored"]["run"])
    source_run_id = run_data["run_id"]
    replay_id = f"RUN-REPLAY-{uuid.uuid4().hex[:12].upper()}"
    run_data["run_id"] = replay_id
    source_usage = {
        "input_tokens": int(run_data.get("input_tokens") or 0),
        "output_tokens": int(run_data.get("output_tokens") or 0),
        "duration_ms": int(run_data.get("duration_ms") or 0),
        "provider_call_count": int(run_data.get("provider_call_count") or 0),
    }
    run_data["context"]["execution_mode"] = "cache_replay"
    run_data["context"]["replayed_from_cache_id"] = cache_id
    run_data["context"]["replayed_from_run_id"] = source_run_id
    run_data["context"]["replayed_at"] = time.strftime("%Y-%m-%dT%H:%M:%S%z")
    run_data["context"]["cache_source_model_usage"] = source_usage
    run_data["context"]["external_model_call_performed"] = False
    run_data["run_completeness"] = "cache_replay_not_fresh_analysis"
    run_data["execution_mode"] = "cache_replay"
    run_data["cache_hit"] = True
    run_data["input_tokens"] = 0
    run_data["output_tokens"] = 0
    run_data["duration_ms"] = 0
    run_data["provider_call_count"] = 0
    run_data["model_check"] = {
        "status": "cache_replay",
        "model_id": run_data.get("model_check", {}).get("model_id"),
        "execution_mode": "cache_replay",
        "cache_hit": True,
        "input_tokens": 0,
        "output_tokens": 0,
        "duration_ms": 0,
        "provider_call_count": 0,
        "detail": "本次回放保留原Agent轨迹，但没有重新运行RAG或模型。",
    }
    # Agent 轨迹属于原运行证据链，回放时保留并由 execution_mode 明确区分。
    return RunResponse.model_validate(run_data)


def _set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_run_font(run: Any, name: str = "Microsoft YaHei") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)


def _add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    _set_run_font(run)


def _table_rows(document: Document, entries: list[tuple[str, Any]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for label, value in entries:
        cells = table.add_row().cells
        cells[0].width = Cm(4.2)
        cells[1].width = Cm(11.2)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[0].text = str(label)
        cells[1].text = str(value)
        _set_cell_shading(cells[0], "DCEFF1")
        for index, cell in enumerate(cells):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run)
                    run.font.size = Pt(8.5)
                    if index == 0:
                        run.bold = True


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        _set_run_font(run)
        run.font.color.rgb = RGBColor(12, 71, 82)


def build_report(workspace_root: Path, stored: StoredRunResponse, *, demo_preview: bool = False) -> Path:
    review = stored.human_review
    if review is None or review.status == "未复核" or not review.export_approved:
        if not demo_preview:
            raise ValueError("导出前必须由人工复核并明确允许导出。")
        review = HumanReviewRequest(
            status="未复核",
            reviewer="竞赛演示模式",
            note="这是用于展示系统思路的自动演示报告，尚未经真人专业复核。",
            reviewed_at=None,
            export_approved=False,
            reviewer_type="automation",
        )
    if stored.run.schema_version != "run_output_v2" or stored.run.context.get("run_schema_version") != "run_output_v2":
        raise ValueError("旧 run_v1 记录只读兼容，不能重新包装成 report_v2。")

    namespace = re.sub(r"[^A-Za-z0-9_-]", "", os.getenv("AUDITTRACE_RUNTIME_NAMESPACE", ""))
    is_automation = demo_preview or review.reviewer_type == "automation" or bool(namespace)
    report_dir = _runtime_base(workspace_root) / "reports"
    report_dir.mkdir(parents=True, exist_ok=True)
    path = report_dir / f"{stored.run.run_id}_预审风险备忘录_{REPORT_VERSION}.docx"
    document = Document()
    document.core_properties.title = "审迹智链预审风险备忘录"
    document.core_properties.comments = AI_GENERATED_CONTENT_NOTICE
    section = document.sections[0]
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("审迹智链 · 预审风险备忘录")
    _set_run_font(title_run)
    title_run.bold = True
    title_run.font.size = Pt(17)
    title_run.font.color.rgb = RGBColor(8, 87, 104)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"{REPORT_VERSION}｜AI生成内容｜审计计划阶段")
    _set_run_font(subtitle_run)
    subtitle_run.bold = True
    subtitle_run.font.color.rgb = RGBColor(22, 124, 142)

    warning = document.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warning_text = (
        "【自动化测试报告｜禁止作为正式备忘录发送】"
        if is_automation
        else f"【{AI_GENERATED_CONTENT_NOTICE}】"
    )
    warning_run = warning.add_run(warning_text)
    _set_run_font(warning_run)
    warning_run.bold = True
    warning_run.font.color.rgb = RGBColor(170, 76, 24)
    if is_automation:
        ai_notice = document.add_paragraph()
        ai_notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ai_notice_run = ai_notice.add_run(f"【{AI_GENERATED_CONTENT_NOTICE}】")
        _set_run_font(ai_notice_run)
        ai_notice_run.bold = True
        ai_notice_run.font.color.rgb = RGBColor(170, 76, 24)
    if demo_preview:
        preview = document.add_paragraph()
        preview.alignment = WD_ALIGN_PARAGRAPH.CENTER
        preview_run = preview.add_run("【竞赛演示报告｜未登录｜未经真人专业复核】")
        _set_run_font(preview_run)
        preview_run.bold = True
        preview_run.font.color.rgb = RGBColor(176, 35, 35)
    if stored.run.run_completeness != "complete_full_analysis":
        incomplete = document.add_paragraph()
        incomplete.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = incomplete.add_run(f"【不完整运行：{stored.run.run_completeness}】")
        _set_run_font(run)
        run.bold = True
        run.font.color.rgb = RGBColor(176, 35, 35)

    document.add_paragraph(
        "边界：本报告不构成舞弊认定、重大错报认定、审计意见或投资建议；来源真实性、专业口径、正式认定和发布均由人负责。"
    )
    context = stored.run.context
    _add_heading(document, "一、运行身份与三层状态", 1)
    _table_rows(
        document,
        [
            ("公司", context.get("company_name", "")),
            ("案例编号", context.get("case_id", "")),
            ("分析时点 T0", context.get("t0", "")),
            ("运行编号", stored.run.run_id),
            ("运行完整性", stored.run.run_completeness),
            ("程序筛查", stored.run.screening_status),
            ("AI建议", stored.run.ai_recommendation),
            ("人工处理", review.status),
            ("复核人 / 时间", f"{review.reviewer or '未填写'} / {review.reviewed_at or '未填写'}"),
            ("资料快照", context.get("source_snapshot_id", "")),
            ("版本链", f"工程 {stored.run.engine_version}｜R1 {context.get('r1_version')}｜提示词 {context.get('agent_prompt_version')}｜{REPORT_VERSION}"),
        ],
    )

    _add_heading(document, "二、程序计算与AI待核查草稿", 1)
    for result in stored.run.rule_results:
        _add_heading(document, f"{result.rule_id}｜程序筛查 {result.screening_status or result.status}", 2)
        card = result.risk_card or {}
        document.add_paragraph(str(card.get("title", "未形成程序筛查摘要")))
        if card.get("observation"):
            document.add_paragraph(str(card["observation"]))
        document.add_paragraph("计算指标：")
        for key, value in result.metrics.items():
            _add_bullet(document, f"{key}: {value}")
        document.add_paragraph(f"AI建议：{result.ai_recommendation}")
        if result.ai_draft:
            document.add_paragraph(AI_GENERATED_CONTENT_NOTICE)
            document.add_paragraph(str(result.ai_draft.get("draft_title") or "AI待核查草稿"))
            document.add_paragraph(str(result.ai_draft.get("draft_observation") or ""))
            for claim in result.ai_draft.get("claims", []):
                _add_bullet(
                    document,
                    f"{claim.get('text')}｜证据 {', '.join(claim.get('evidence_ids', []))}｜{claim.get('support_status')}",
                )
            for explanation in result.ai_draft.get("normal_explanations", []):
                _add_bullet(
                    document,
                    f"正常解释候选：{explanation.get('text')}｜{explanation.get('support_status')}｜证据 {', '.join(explanation.get('evidence_ids', [])) or '无'}",
                )
            gaps = result.ai_draft.get("data_gaps") or []
            if gaps:
                document.add_paragraph("资料缺口：" + "；".join(map(str, gaps)))
            requested = result.ai_draft.get("requested_materials") or []
            if requested:
                document.add_paragraph("待索取资料：" + "；".join(map(str, requested)))
        else:
            document.add_paragraph("未形成通过硬校验的AI草稿；不得把程序筛查摘要当作AI完整分析。")

    _add_heading(document, "三、RAG检索记录与候选原文", 1)
    if not stored.run.retrievals:
        document.add_paragraph("本次没有RAG检索记录（无候选或运行未完成）。")
    for retrieval in stored.run.retrievals:
        question = retrieval.get("question") or {}
        document.add_paragraph(
            f"{retrieval.get('retrieval_id')}｜{question.get('question_id', '自由检索')}｜{retrieval.get('status')}｜{retrieval.get('retrieval_version')}"
        )
        for item in retrieval.get("results", []):
            _add_bullet(
                document,
                f"{item.get('evidence_id')}｜{item.get('document_id')}｜PDF第{item.get('pdf_page')}页｜{item.get('source_locator')}｜{str(item.get('excerpt') or '')[:260]}",
            )
        if retrieval.get("status") == "no_hit":
            document.add_paragraph(str((retrieval.get("evidence_gap") or {}).get("message", "本次未命中。")))

    _add_heading(document, "四、三Agent轨迹", 1)
    for result in stored.run.rule_results:
        document.add_paragraph(f"{result.rule_id}：")
        for step in result.agent_steps:
            _add_bullet(
                document,
                f"{step.role}｜{step.status}｜模型 {step.model_id or '未调用'}｜提示词 {step.prompt_version or '不适用'}｜{step.detail}",
            )

    _add_heading(document, "五、字段、补充资料与来源回查", 1)
    for source in stored.run.evidence_bundle.get("field_evidence", []):
        _add_bullet(
            document,
            f"{source.get('evidence_id')}｜{source.get('document_id')}｜{source.get('source_file')}｜PDF第{source.get('pdf_page')}页｜{source.get('locator')}｜SHA-256 {source.get('file_sha256', '')}",
        )
    supplements = stored.run.evidence_bundle.get("supplement_evidence", [])
    if supplements:
        document.add_paragraph("补充资料（as_of_date 与原 T0 分开）：")
        for item in supplements:
            _add_bullet(
                document,
                f"{item.get('evidence_id')}｜{item.get('field_label')}｜as_of {item.get('as_of_date')}｜{item.get('support_status')}",
            )

    _add_heading(document, "六、人工复核说明", 1)
    document.add_paragraph(review.note or "未填写补充说明。")
    document.add_paragraph(
        f"AI生成内容声明：{AI_GENERATED_CONTENT_NOTICE} AI只参与受证据编号约束的待核查语义草稿。程序计算、来源哈希、T0过滤、人工处理及正式发布均不由模型决定。"
    )
    document.save(path)
    return path
