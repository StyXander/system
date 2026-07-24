"""审迹智链 W2 三份专业底稿的审查修订工具。

使用方式（请通过本项目附带的 Python 运行）：
    python 审查修订_W2规则底稿.py --inspect
    python 审查修订_W2规则底稿.py --apply

本脚本只处理本目录内三份 2026-07-24 提交件。修订原则：
1. 保留规则草案的专业思路；
2. 标出“规则草案”与“当前工程实现”的边界；
3. 统一字段名为网页/后端正在使用的 *_previous 口径；
4. 不把任何自动结果写成审计结论。
"""

from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = Path(__file__).resolve().parent
FILES = {
    "r1": BASE / "R1 应收款项与营业收入增速匹配性预筛查规则 v0.3(1).docx",
    "r2": BASE / "R2 收入现金质量核查规则 v2(1).docx",
    "map": BASE / "销售与收款循环映射表v2.docx",
}


def dump_document(path: Path) -> None:
    """输出可复核的正文与表格文字，不写入文件。"""
    doc = Document(path)
    print(f"\n===== {path.name} =====")
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:
            print(f"P{index}: {text}")
    for table_index, table in enumerate(doc.tables):
        print(f"-- TABLE {table_index}: {len(table.rows)} x {len(table.columns)} --")
        for row_index, row in enumerate(table.rows):
            values = [cell.text.replace("\n", " / ").strip() for cell in row.cells]
            print(f"R{row_index}: " + " | ".join(values))


def replace_paragraph(paragraph, text: str) -> None:
    """以整段替换处理审查结论，保留段落的原有段落样式。"""
    paragraph.text = text


def replace_when_contains(doc: Document, needle: str, replacement: str, *, limit: int | None = None) -> int:
    """在正文中作最小定位替换；找不到会在最终结构检查中暴露。"""
    changed = 0
    for paragraph in doc.paragraphs:
        if needle in paragraph.text and (limit is None or changed < limit):
            replace_paragraph(paragraph, replacement)
            changed += 1
    return changed


def set_cell(cell, text: str, *, bold: bool = False, color: str | None = None, size: float = 8.4) -> None:
    """设置单元格正文，并保持映射表的紧凑、可读格式。"""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.04
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = 120) -> None:
    """让 Word 的 table/grid/cell 宽度一致，避免自动布局带来的表格漂移。"""
    total = sum(widths_dxa)
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths_dxa):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def normalize_existing_table_geometry(doc: Document) -> None:
    """不改原表结构，只把现有列宽、表宽和单元格宽度统一为可复核 DXA。"""
    for table in doc.tables:
        widths = []
        for grid_col in table._tbl.tblGrid.gridCol_lst:
            value = grid_col.get(qn("w:w"))
            widths.append(int(value) if value and value.isdigit() else 0)
        if not widths or len(widths) != len(table.columns) or not all(widths):
            widths = [8280 // len(table.columns)] * len(table.columns)
            widths[-1] += 8280 - sum(widths)
        set_table_geometry(table, widths)


def insert_paragraph_after(paragraph, text: str, *, italic: bool = False, color: str | None = None) -> None:
    """在原位置插入审查说明，避免把修订理由全部堆到文末。"""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p.getparent().remove(new_paragraph._p)
    new_paragraph._p = new_p
    run = new_paragraph.add_run(text)
    run.italic = italic
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    new_paragraph.paragraph_format.space_before = Pt(3)
    new_paragraph.paragraph_format.space_after = Pt(3)


def revise_r1() -> None:
    path = FILES["r1"]
    doc = Document(path)

    # 统一网页、后端和规则草案三者使用的字段名；先处理较长的 prior_prior，避免重复替换。
    field_replacements = [
        ("ar_prior_prior", "ar_previous_previous"),
        ("revenue_prior", "revenue_previous"),
        ("ar_prior", "ar_previous"),
        ("bad_debt_current/prior", "bad_debt_current/bad_debt_previous"),
        ("contract_asset_current/prior", "contract_asset_current/contract_asset_previous"),
    ]
    for paragraph in doc.paragraphs:
        text = paragraph.text
        for before, after in field_replacements:
            text = text.replace(before, after)
        if text != paragraph.text:
            replace_paragraph(paragraph, text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                for before, after in field_replacements:
                    text = text.replace(before, after)
                if text != cell.text:
                    cell.text = text

    replace_when_contains(doc, "R1 应收款项与营业收入增速匹配性预筛查规则 v0.3", "R1 应收款项与营业收入增速匹配性预筛查规则 v0.3.1（审查修订稿）")
    replace_when_contains(doc, "R1 应收 - 收入背离核查规则 v0.3", "R1 应收 - 收入背离核查规则 v0.3.1（审查修订稿）")
    replace_when_contains(doc, "规则版本：v0.3", "规则版本：v0.3.1（审查修订稿，待专业确认）")
    replace_when_contains(
        doc,
        "专业依据：《中国注册会计师审计准则第 1313 号",
        "专业依据：以《中国注册会计师审计准则第 1211 号——重大错报风险的识别和评估》与第 1313 号——分析程序为准则基础。本规则将应收—收入背离作为分析性程序下的风险识别线索，不将其作为错报、舞弊或收入真实性结论。",
    )
    boundary_count = replace_when_contains(
        doc,
        "本规则仅分析应收款项与营业收入的增速匹配性",
        "本规则仅分析应收款项与营业收入的增速匹配性，不覆盖收入确认政策、成本结转、关联方交易等其他收入相关风险；未触发本规则不代表不存在其他收入确认相关风险。当前工程版只读取 revenue_current、revenue_previous、ar_current、ar_previous 四个“应收账款”字段做方向比较；本稿中的应收票据、合同资产、阈值分级与跨规则联动均属于待专业确认、待代码对齐的规则草案。",
    )
    if boundary_count:
        target = next(p for p in doc.paragraphs if "当前工程版只读取 revenue_current" in p.text)
        insert_paragraph_after(
            target,
            "口径提示：营业收入是不含税的期间流量，应收款项可能为含税的期末余额；本规则只观察增速趋势，不能把两者解释为一一对应的金额或回款比例。",
            italic=True,
            color="5B6F82",
        )
    replace_when_contains(
        doc,
        "上年应收款项余额绝对额过小",
        "上年应收款项余额绝对额过小（草案默认＜100 万元）或占上年营业收入比例过低（草案默认＜5%）时，标记为“基数过小，不可比”，不触发风险判定；参数须经专业确认后才可进入工程配置。",
    )
    replace_when_contains(doc, "通用阈值", "草案通用阈值（未接入当前工程）")
    replace_when_contains(doc, "增速差值＞15 个百分点：", "草案参数：增速差值＞15 个百分点时，建议人工关注；不由当前工程自动分级。")
    replace_when_contains(doc, "增速差值＞30 个百分点：", "草案参数：增速差值＞30 个百分点时，建议人工重点复核；不由当前工程自动升级。")
    replace_when_contains(
        doc,
        "工程建筑、大型设备制造、公共事业等长账期行业",
        "行业阈值适配为后续专业配置项。v0.3.1 首版不按行业自动适配，不以“通用行业经验”替代经专业确认的行业参数。",
    )
    replace_when_contains(doc, "无对应行业配置时，输出结果标注", "未提供经专业确认的行业参数时，不输出行业阈值结论，只保留“待人工结合行业特征判断”的说明。")
    replace_when_contains(
        doc,
        "收入负增长 + 应收正增长为强异常信号",
        "收入负增长且应收正增长属于需要人工优先复核的草案情形，不由当前系统直接认定为重点核查或自动升级。",
    )
    replace_when_contains(doc, "（二）三级风险等级（与 R2 规则完全对齐）", "（二）三级风险等级（规则草案；当前页面未按此分级）")
    replace_when_contains(doc, "应收 - 收入增速差值＞30 个百分点；", "草案条件：应收 - 收入增速差值＞30 个百分点；")
    replace_when_contains(doc, "收入负增长且应收正增长；", "草案条件：收入负增长且应收正增长；")
    replace_when_contains(doc, "连续两期触发", "连续两期满足草案关注条件时，作为人工复核参考，不由系统自动升级。")
    replace_when_contains(doc, "单期触发待关注，且联动其他规则同步触发。", "单期满足草案关注条件且已接入的其他规则独立形成候选时，可提交人工评估是否合并线索；不得自动升级。")
    replace_when_contains(
        doc,
        "系统仅基于公开数据自动计算指标并判定触发状态",
        "当前工程版仅基于公开数据完成四字段方向比较并输出候选状态；v0.3.1 的阈值分级尚未接入。系统不得自动降级或关闭线索；所有正常解释仅作为参考标注，最终是否保留及后续处理由审计人员人工确认。",
    )
    replace_when_contains(
        doc,
        "AGENT 调用规则：仅当规则触发（待关注 / 待重点核查）时",
        "AGENT 调用规则：当前工程版仅在四字段方向比较形成候选时调用模型；v0.3.1 的阈值等级完成代码与测试前，不得据此宣称模型已按“待关注/待重点核查”触发。",
    )
    replace_when_contains(
        doc,
        "R1+R2 双触发：应收增速背离",
        "R1+R2 联动：仅当两条规则均完成独立字段、来源与可比性校验并形成候选时，才提交人工评估是否合并为一项线索。当前工程无法识别 OCF 变动“主因”，不得据此自动升级。",
    )
    replace_when_contains(
        doc,
        "R1+R8 联动：年末 / 四季度收入集中",
        "R1+R8 联动为后续规则库设计：R8 当前未接入计算，不得因勾选 R8 自动改变 R1 结果；取得分季度收入、退货折让和期后资料后由人工判断是否开展截止性核查。",
    )
    replace_when_contains(
        doc,
        "R1+R4 联动：前五大客户集中度提升",
        "R1+R4 联动为后续规则库设计：R4 当前未接入计算，客户集中度与公开特征资料须先核验；不得自动扩大函证范围。",
    )
    replace_when_contains(
        doc,
        "仅 R1 触发、R2 未触发：提示",
        "仅 R1 形成候选、R2 未形成候选时：仅说明在当前已接入字段及可比条件下未出现 R2 方向候选；仍需人工结合应收可回收性、坏账计提与现金流构成判断，不能推导“收入现金质量无异常”。",
    )
    replace_when_contains(doc, "复核结论统一设置 5 类选项：未复核、无异常关闭", "复核结论统一设置 5 类选项：未复核、人工确认不保留")
    replace_when_contains(
        doc,
        "已复核：经审计专业成员审核",
        "专业复核完成：经审计专业成员审核，准则依据准确、公式逻辑无误、禁用表述控制到位、口径无歧义；",
    )
    replace_when_contains(
        doc,
        "已实现：后端编码完成",
        "工程对齐完成：后端编码完成并能按本稿字段、可比性与阈值逻辑输出结果，与 Excel 人工复算偏差≤0.01%。本条仅为验收条件，当前未完成；",
    )
    replace_when_contains(
        doc,
        "已测试：通过≥5 组测试用例",
        "测试完成：通过≥5 组测试用例（正常、草案关注、草案重点复核、数据缺失、基数不可比），输出与预期一致；",
    )

    # 必选字段表：把“扩展口径”与当前实现拆开。
    core = doc.tables[0]
    core.cell(2, 1).text = "revenue_previous"
    core.cell(3, 2).text = "合并资产负债表“应收账款”账面余额。当前工程版只支持该字段；应收票据和合同资产为专业规则草案扩展项，不得当作当前已运行字段。"
    core.cell(4, 1).text = "ar_previous"
    optional = doc.tables[2]
    optional.cell(1, 1).text = "ar_previous_previous"
    optional.cell(2, 1).text = "bad_debt_current/bad_debt_previous"
    optional.cell(3, 1).text = "contract_asset_current/contract_asset_previous"
    params = doc.tables[4]
    extra = params.add_row().cells
    extra[0].text = "是否纳入合同资产"
    extra[1].text = "false（草案默认）"
    extra[2].text = "仅适用于未来扩展口径；当前工程未接入。"
    for row in params.rows[1:6]:
        row.cells[2].text += "（草案参数，待专业确认）"

    doc.add_paragraph("v0.3.1，2026 年 7 月 24 日：审查修订。统一字段 ID 为网页/后端的 *_previous 命名；明确当前工程仅为四字段方向比较；将阈值、行业适配、扩展应收口径与跨规则联动标注为待专业确认、待代码对齐的规则草案；补充含税应收与不含税收入不可直接作金额对应的口径提示。")
    doc.save(path)


def revise_r2() -> None:
    path = FILES["r2"]
    doc = Document(path)

    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("R2 收入现金质量核查规则"):
            replace_paragraph(paragraph, "R2 收入现金质量核查规则 v0.2.1（审查修订稿）")
    replace_when_contains(doc, "规则版本：v2", "规则版本：v0.2.1（审查修订稿，待专业确认）")
    replace_when_contains(
        doc,
        "专业依据：《中国注册会计师审计准则第 1313 号",
        "专业依据：以《中国注册会计师审计准则第 1211 号——重大错报风险的识别和评估》与第 1313 号——分析程序为准则基础。R2 只用于识别收入、经营现金流和收现质量之间需要解释的线索，不构成错报、舞弊或收入真实性结论。",
    )
    boundary_count = replace_when_contains(
        doc,
        "本规则仅分析收入与经营活动现金流量净额",
        "本规则分析收入与经营活动现金流量、销售收现和净利润现金含量之间的待解释线索，不覆盖全部现金流风险。当前工程版只读取 revenue_current、revenue_previous、operating_cash_flow_current、operating_cash_flow_previous、net_profit_current，按“收入正增长下的经营现金流方向比较”运行；销售收现率、OCF 构成分析、行业阈值与本稿完整分级尚待接入和验收。",
    )
    if boundary_count:
        target = next(p for p in doc.paragraphs if "当前工程版只读取 revenue_current" in p.text)
        insert_paragraph_after(
            target,
            "可比性提示：经营现金流跨期变号、上年 OCF 基数很小或来源口径不一致时，不展示经营现金流同比百分比或增速差，应标记为“同比不宜比较”，改看金额、收现资料与构成说明。",
            italic=True,
            color="5B6F82",
        )
    replace_when_contains(
        doc,
        "本指标受应收票据背书转让、销售退回、合同负债转回等因素影响",
        "销售收现率受应收票据背书转让、销售退回、合同负债转回等因素影响。销售商品、提供劳务收到的现金通常含增值税，并受应收、预收变动影响；营业收入通常不含税，因此该比率可能超过 1，不能被视为收入的机械现金对应关系。",
    )
    replace_when_contains(
        doc,
        "通用阈值：收现率 0.8 为异常关注阈值",
        "草案通用阈值：收现率 0.8 仅作为后续测试候选，不是当前工程触发条件，也不是已验证的行业标准。",
    )
    replace_when_contains(
        doc,
        "工程、建筑、大型设备制造等长账期行业",
        "行业差异需在取得经专业确认的行业口径后另行配置。v0.2.1 首版不按行业自动适配，也不以笼统行业经验替代参数。",
    )
    replace_when_contains(doc, "无对应行业配置时，输出", "未提供经专业确认的行业参数时，只标记为“待人工结合行业与合同结算条件判断”，不输出行业阈值结论。")
    replace_when_contains(doc, "仅当净利润为正时，展示净利润现金含量具体比率数值；", "仅当净利润为正时，展示净利润现金含量具体比率数值；净利润为零或负数时不计算、不展示该比例。")
    replace_when_contains(
        doc,
        "净利润为负时，重点关注",
        "净利润为负且 OCF 为非负时，显示“亏损期 OCF 非负，需人工结合利润构成与营运资本变动解释”；净利润为负且 OCF 也为负时，显示“亏损期 OCF 为负，需补充收现和构成资料”。两种情形均不以该指标自动升级 R2。",
    )
    replace_when_contains(doc, "净利润现金含量异常阈值：OCF / 净利润 < 0.5", "草案参数：净利润为正时，OCF / 净利润 < 0.5 可作为待人工关注的候选条件；当前工程不以此触发。")
    replace_when_contains(
        doc,
        "上年 OCF 绝对值 ≥ 上年营业收入 × 3%",
        "草案可比前提：上年 OCF 绝对值 ≥ 上年营业收入 × 3%，且本期与上期 OCF 不跨期变号。未满足时标记“同比不宜比较”；当前工程正按此方向对齐。",
    )
    replace_when_contains(doc, "上年营业收入 ≥ 可比基数（默认 1000 万元，可配置）；", "上年营业收入 ≥ 草案可比基数（默认 1000 万元，待专业确认）；")
    replace_when_contains(doc, "收入增速 - OCF 增速 > 阈值（默认 20 个百分点，可配置）时", "草案参数：收入增速 - OCF 增速 > 阈值（默认 20 个百分点，待专业确认）时，才可作为趋势背离候选；跨期变号或基数过小时不适用。")
    replace_when_contains(
        doc,
        "强收入关联：若差异主要由经营性应收项目增加、合同负债减少导致",
        "较强收入端关联（需有 OCF 构成证据）：若差异主要由经营性应收项目增加、合同负债减少导致，可作为与 R1 联动的人工复核参考；不得由系统自动提高风险等级。",
    )
    replace_when_contains(
        doc,
        "弱收入关联：若差异主要由存货备货增加、经营性应付项目减少导致",
        "较弱收入端关联（需有 OCF 构成证据）：若差异主要由存货备货增加、经营性应付项目减少导致，提示同时关注采购、生产或付款循环；不自动降低、关闭或移除 R2 线索。",
    )
    replace_when_contains(
        doc,
        "若差异主要由折旧摊销、资产减值、投资收益等非现金 / 非经营项目导致",
        "若差异主要由折旧摊销、资产减值、投资收益等非现金或非经营项目导致，作为解释候选并要求补充构成明细；是否保留线索由人工确认。",
    )
    replace_when_contains(doc, "（二）三级风险等级（与 R1 规则对齐）", "（二）三级风险等级（规则草案；当前页面未按此分级）")
    replace_when_contains(
        doc,
        "双触发且 OCF 背离主因为应收增加：强关联信号，收入回款真实性风险优先级最高，合并输出风险线索。",
        "R1、R2 同时形成候选且经构成资料证实 OCF 变动与应收端相关时，可提交人工评估是否合并为一项“收入与回款匹配性待核查线索”；不得自动升级，也不得使用“真实性风险最高”等定性表述。",
    )
    replace_when_contains(
        doc,
        "仅 R2 触发、R1 未触发且 OCF 背离主因为存货 / 应付：提示",
        "仅 R2 形成候选、R1 未形成候选且经构成资料显示 OCF 变动主要来自存货或应付时：提示该说明需要人工核验并可建议联动存货、付款循环；不自动降低 R2 或断言收入端风险较低。",
    )
    replace_when_contains(
        doc,
        "仅 R1 触发、R2 未触发：提示",
        "仅 R1 形成候选、R2 未形成候选时：仅说明当前 R2 已接入字段及可比条件下未形成方向候选；不得据此推导收入现金质量不存在异常。",
    )
    replace_when_contains(
        doc,
        "已复核：经审计专业成员审核",
        "专业复核完成：经审计专业成员审核，准则依据准确、公式逻辑无误、禁用表述控制到位、口径无歧义；",
    )
    replace_when_contains(
        doc,
        "已实现：后端编码完成",
        "工程对齐完成：后端编码完成并能按本稿字段、可比性与阈值逻辑输出结果，与 Excel 人工复算偏差≤0.01%。本条仅为验收条件，当前未完成；",
    )
    replace_when_contains(
        doc,
        "已测试：通过≥5 组测试用例",
        "测试完成：通过≥5 组测试用例（正常、草案触发、未触发、数据缺失、不可比），输出与预期一致；",
    )

    field_table = doc.tables[0]
    field_table.cell(1, 1).text = "合并利润表“营业收入”项目，通常为不含税口径；与销售收现只作趋势与解释性比较，不直接作价税一致性或一一对应判断。"
    field_table.cell(3, 1).text = "直接取自合并现金流量表对应行次，通常含增值税并受应收、预收变动影响；本字段为 v0.2.1 草案收现率所需字段，当前工程未接入。"
    params = doc.tables[2]
    for row in params.rows[1:]:
        row.cells[2].text += "（草案参数，待专业确认；当前工程未以此触发）"
    extra = params.add_row().cells
    extra[0].text = "行业阈值适配"
    extra[1].text = "未启用"
    extra[2].text = "v0.2.1 首版不自动按行业适配；需专业确认后另行配置。"

    doc.add_paragraph("v0.2.1，2026 年 7 月 24 日：审查修订。补足 1211/1313 的规则边界；明确销售收现含税与收入不含税不能直接比值定性；规定净利润为负时不展示现金含量比率；经营现金流跨期变号或基数过小时标记“同比不宜比较”；取消 OCF 构成对风险线索的自动升降级表述；明确当前工程仍为方向比较版。")
    doc.save(path)


def repair_interrupted_inline_notes() -> None:
    """首轮保存中断后，只补入缺失的定位说明，避免重复重跑整份修订。"""
    r1_path = FILES["r1"]
    r1 = Document(r1_path)
    r1_paragraphs = list(r1.paragraphs)
    # 当前稿的规则正文从 P30 开始，P69 是“30 个百分点”的参数行。
    # 这一定点修正只处理本次已中断保存的同一份稿件，避免干扰文首摘要。
    if len(r1_paragraphs) <= 69:
        raise RuntimeError("未找到 R1 草案阈值位置。")
    replace_paragraph(r1_paragraphs[69], "草案参数：增速差值＞30 个百分点时，建议人工重点复核；不由当前工程自动升级。")
    replace_when_contains(r1, "极端场景强制升级", "极端场景（人工优先复核，不自动升级）")
    replace_when_contains(r1, "上年应收款项账面余额 ≥ 100 万元", "草案可比前提：上年应收款项账面余额 ≥ 100 万元（待专业确认），且占上年营业收入比例 ≥ 5%（待专业确认）；")
    replace_when_contains(r1, "上年营业收入 ≥ 1000 万元", "草案可比前提：上年营业收入 ≥ 1000 万元（待专业确认）；")
    normalize_existing_table_geometry(r1)
    r1.save(r1_path)

    r2_path = FILES["r2"]
    r2 = Document(r2_path)
    replace_when_contains(r2, "规则版本：v0.2", "规则版本：v0.2.1（审查修订稿，待专业确认）")
    all_r2 = "\n".join(p.text for p in r2.paragraphs)
    if "当前工程版只读取 revenue_current" not in all_r2:
        count = replace_when_contains(
            r2,
            "经营活动产生的现金流量净额（以下简称OCF）",
            "本规则分析收入与经营活动现金流量、销售收现和净利润现金含量之间的待解释线索，不覆盖全部现金流风险。当前工程版只读取 revenue_current、revenue_previous、operating_cash_flow_current、operating_cash_flow_previous、net_profit_current，按“收入正增长下的经营现金流方向比较”运行；销售收现率、OCF 构成分析、行业阈值与本稿完整分级尚待接入和验收。",
            limit=1,
        )
        if not count:
            raise RuntimeError("未找到可写入 R2 当前工程边界的位置。")
        target = next(p for p in r2.paragraphs if "当前工程版只读取 revenue_current" in p.text)
        insert_paragraph_after(
            target,
            "可比性提示：经营现金流跨期变号、上年 OCF 基数很小或来源口径不一致时，不展示经营现金流同比百分比或增速差，应标记为“同比不宜比较”，改看金额、收现资料与构成说明。",
            italic=True,
            color="5B6F82",
        )
    replace_when_contains(
        r2,
        "本指标受应收票据背书转让",
        "销售收现率受应收票据背书转让、销售退回、合同负债结转等因素影响。销售商品、提供劳务收到的现金通常含增值税，并受应收、预收变动影响；营业收入通常不含税，因此该比率可能超过 1，不能被视为收入的机械现金对应关系。",
    )
    replace_when_contains(
        r2,
        "通用阈值：收现率＜0.8",
        "草案通用阈值：收现率＜0.8 仅作为后续测试候选，不是当前工程触发条件，也不是已验证的行业标准。",
    )
    replace_when_contains(
        r2,
        "净利润为负时，仅描述方向组合",
        "净利润为负且 OCF 为非负时，显示“亏损期 OCF 非负，需人工结合利润构成与营运资本变动解释”；净利润为负且 OCF 也为负时，显示“亏损期 OCF 为负，需补充收现和构成资料”。两种情形均不以该指标自动升级 R2。",
    )
    replace_when_contains(
        r2,
        "净利润为正且净利润现金含量＜0.5",
        "草案参数：净利润为正且净利润现金含量＜0.5 可作为待人工关注的候选条件；当前工程不以此触发。",
    )
    replace_when_contains(r2, "本年营业收入 ≥ 可比基数", "上年营业收入 ≥ 草案可比基数（默认 1000 万元，待专业确认）")
    replace_when_contains(
        r2,
        "当「收入增速 - OCF 增速」",
        "草案参数：收入增速 - OCF 增速超过默认 20 个百分点（待专业确认）且满足可比前提时，才可作为趋势背离候选；跨期变号或基数过小时不适用。",
    )
    replace_when_contains(
        r2,
        "无收入关联：若差异主要由折旧摊销",
        "若差异主要由折旧摊销、资产减值、投资收益等非现金或非经营项目导致，作为解释候选并要求补充构成明细；是否保留线索由人工确认。",
    )
    normalize_existing_table_geometry(r2)
    r2.save(r2_path)


def remove_all_body_tables(doc: Document) -> None:
    for table in list(doc.tables):
        table._element.getparent().remove(table._element)


def revise_map() -> None:
    """旧表只有宽泛异常描述，重建为可直接支持 W2 规则验收的循环映射。"""
    path = FILES["map"]
    doc = Document(path)
    remove_all_body_tables(doc)

    # 保留原文档，但将两行正文改为版本和使用边界，避免“AI 生成”被误读为资料来源。
    replace_paragraph(doc.paragraphs[0], "销售与收款循环映射表 v0.3（对齐 R1 v0.3.1 / R2 v0.2.1 草案）")
    replace_paragraph(doc.paragraphs[1], "用途：业务承接、续聘复核和审计计划阶段的公开资料预筛；只形成待核查线索、资料缺口与建议索取资料，不形成审计结论。")
    p = doc.add_paragraph("当前状态：R1、R2 工程版可进行来源绑定和方向计算；本表中的阈值分级、扩展字段、OCF 构成与 R3—R8 均为待专业确认或待接入内容，不能表述为已运行。")
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    p = doc.add_paragraph("修订说明：将原表的宽泛“异常信号”改为“公开资料线索 + 规则状态 + 正常解释 + 待补资料/人工后续”五列，以便与 R1/R2 规则底稿和主网页的运行边界一致。")
    p.paragraph_format.space_after = Pt(8)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.35)

    headers = ["业务环节", "规则及当前状态", "公开资料线索", "正常解释候选", "待补资料 / 人工后续"]
    rows = [
        [
            "客户准入与信用政策",
            "R4 候选规则\n未接入计算",
            "前五大客户集中度、信用期或客户结构明显变化。",
            "拓展头部客户、退出低毛利客户、行业账期整体后移。",
            "客户清单、信用政策变更、主要合同账期；由项目组确认是否扩大客户背景核查。",
        ],
        [
            "合同、发货与验收",
            "R8 候选规则\n未接入计算",
            "年末/第四季度收入集中，或发货、验收、结算时点可能跨期。",
            "季节性订单、年末集中交付、合同验收周期差异。",
            "分季度收入、期末大额合同、发货/验收/退货资料；人工判断是否安排截止性核查。",
        ],
        [
            "收入确认与开票",
            "R1 v0.3.1 草案\n当前工程：四字段方向比较",
            "应收账款增速高于营业收入增速；当前工程仅比 revenue/ar 的本年、上年四个字段。",
            "信用政策放宽、新增客户、结算节奏变化、行业账期拉长。",
            "应收账款账龄、信用条款、期后回款、主要合同；阈值分级和应收票据/合同资产扩展待接入。",
        ],
        [
            "应收与回款管理",
            "R1 后续核查\n人工确认",
            "期末应收规模、账龄迁徙、期后回款与收入变化需要解释。",
            "年末集中发货、票据背书/贴现、催收节奏变化、合并范围变化。",
            "账龄分段、期后回款、票据背书/贴现、客户往来对账；由人工判断可回收性和后续程序。",
        ],
        [
            "销售收现与经营现金流",
            "R2 v0.2.1 草案\n当前工程：方向比较版",
            "收入增长而 OCF 未同步改善；OCF 跨期变号或基数很小时不展示同比百分比。",
            "存货备货、应付集中支付、税费缴纳、预收结转、票据背书等。",
            "销售收现、季度现金流、经营性应收应付和 OCF 构成明细；销售收现率、净利润现金含量与构成分析尚待接入。",
        ],
        [
            "产品结构、成本与毛利",
            "R3 / R6 候选规则\n未接入计算",
            "产销量与收入、毛利率与同业趋势出现需解释的背离。",
            "产品结构升级、外协补充产能、成本管控、原材料价格变化。",
            "分产品产销量、成本结构、同行同口径数据；先完成来源与可比性复核。",
        ],
        [
            "披露一致性与期后事项",
            "R7 / R8 候选规则\n未接入计算",
            "管理层表述与数据趋势不一致，或期末交易集中且缺少解释。",
            "披露统计口径差异、季节性、年末重大订单集中交付。",
            "年报披露原文、分季度数据、期后退货与结算资料；由人工判断是否保留披露一致性线索。",
        ],
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        # 个别队员模板没有英文内置表格样式；下面仍会写入完整边框、底纹与几何。
        pass
    set_repeat_table_header(table.rows[0])
    widths = [1450, 2100, 2150, 2400, 2700]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell(cell, text, bold=True, color="FFFFFF", size=8.8)
        set_cell_shading(cell, "1E4E69")
    for row_values in rows:
        cells = table.add_row().cells
        for index, (cell, text) in enumerate(zip(cells, row_values)):
            set_cell(cell, text, bold=(index == 0), size=8.15)
            if index == 1:
                set_cell_shading(cell, "EAF2F8")
            elif len(table.rows) % 2 == 0:
                set_cell_shading(cell, "F8FBFD")
    set_table_geometry(table, widths)

    note = doc.add_paragraph("编写说明：本表由团队成员负责专业复核；AI 仅辅助语言组织与格式，不作为资料来源、专业结论或人工复核记录。")
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(0)
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(8.5)
    note.runs[0].font.color.rgb = RGBColor(91, 111, 130)
    doc.save(path)


def structural_checks() -> None:
    """不依赖 Word/LibreOffice 的最低交付检查。"""
    r1 = Document(FILES["r1"])
    r1_all = "\n".join(p.text for p in r1.paragraphs) + "\n" + "\n".join(c.text for t in r1.tables for row in t.rows for c in row.cells)
    assert "revenue_prior" not in r1_all and "ar_prior" not in r1_all, "R1 仍残留旧字段 ID"
    assert "revenue_previous" in r1_all and "ar_previous" in r1_all, "R1 未完成字段 ID 对齐"
    assert "当前工程版只读取" in r1_all, "R1 未标注工程边界"

    r2 = Document(FILES["r2"])
    r2_all = "\n".join(p.text for p in r2.paragraphs) + "\n" + "\n".join(c.text for t in r2.tables for row in t.rows for c in row.cells)
    assert "同比不宜比较" in r2_all, "R2 未写入可比性保护"
    assert "不自动降低、关闭或移除" in r2_all, "R2 仍允许自动降级"
    assert "当前工程版只读取" in r2_all, "R2 未标注工程边界"

    mapping = Document(FILES["map"])
    assert len(mapping.tables) == 1 and len(mapping.tables[0].rows) == 8 and len(mapping.tables[0].columns) == 5, "映射表结构不符合预期"
    mapping_all = "\n".join(c.text for row in mapping.tables[0].rows for c in row.cells)
    assert "R1 v0.3.1 草案" in mapping_all and "R2 v0.2.1 草案" in mapping_all, "映射表未对齐 R1/R2 规则状态"


def apply() -> None:
    # 允许中断后安全重跑，不重复追加版本记录或参数行。
    r1_text = "\n".join(p.text for p in Document(FILES["r1"]).paragraphs)
    r2_text = "\n".join(p.text for p in Document(FILES["r2"]).paragraphs)
    if "v0.3.1（审查修订稿" in r1_text and "v0.2.1（审查修订稿" in r2_text:
        repair_interrupted_inline_notes()
    else:
        revise_r1()
        revise_r2()
        repair_interrupted_inline_notes()
    revise_map()
    structural_checks()
    print("三份 W2 文档已完成审查修订与结构检查。")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--inspect", action="store_true", help="只读取，不修改")
    parser.add_argument("--apply", action="store_true", help="执行三份文档的审查修订")
    parser.add_argument("--map-only", action="store_true", help="只重建销售与收款循环映射表（用于中断恢复）")
    parser.add_argument("--r2-paras", action="store_true", help="仅输出 R2 正文，便于中断时定位")
    args = parser.parse_args()

    if args.inspect or (not args.apply and not args.map_only):
        if args.r2_paras:
            doc = Document(FILES["r2"])
            for index, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    print(f"P{index}: {paragraph.text}")
            return
        for path in FILES.values():
            dump_document(path)
        return

    if args.map_only:
        repair_interrupted_inline_notes()
        revise_map()
        structural_checks()
        print("销售与收款循环映射表已完成审查修订与结构检查。")
        return

    apply()


if __name__ == "__main__":
    main()
