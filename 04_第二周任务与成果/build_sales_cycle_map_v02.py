"""生成 W2 销售与收款循环映射表 v0.2。

本文件把原 v0.1 的段落式说明整理为可核对的业务循环矩阵；不改动原始 Word。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips


BASE_DIR = Path(__file__).resolve().parent
OUTPUT = BASE_DIR / "销售与收款循环映射表_v0.2_对齐V3.docx"

BLUE = "1F4E78"
LIGHT_BLUE = "DDEBF7"
LIGHT_GRAY = "F2F4F7"
LIGHT_GREEN = "E2F0D9"
LIGHT_AMBER = "FFF2CC"
MID_GRAY = "5B6573"


def set_run_font(run, name: str = "Microsoft YaHei", size: float = 9.5, bold: bool = False, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 70, start: int = 90, bottom: int = 70, end: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_cm: list[float]) -> None:
    """同步 tblW、tblGrid 和 tcW；必须在所有行写入后再调用。"""
    widths = [round(width * 567) for width in widths_cm]
    widths[-1] += round(sum(widths_cm) * 567) - sum(widths)
    total_width = sum(widths)
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "90")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Twips(widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def write_cell(cell, text: str, *, header: bool = False, fill: str | None = None, size: float = 8.2) -> None:
    cell.text = ""
    if fill:
        shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=header, color="FFFFFF" if header else None)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, size=12.5 if level == 1 else 10.5, bold=True, color=BLUE)


def add_body(doc: Document, text: str, *, size: float = 9.5, color: str | None = None, bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("审迹智链 AuditTrace ｜ W2 专业底稿")
    set_run_font(header_run, size=8, color=MID_GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("仅用于审计前置阶段的风险预筛与资料索取；不构成审计结论。")
    set_run_font(footer_run, size=8, color=MID_GRAY)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("销售与收款循环映射表")
    set_run_font(title_run, size=19, bold=True, color=BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(9)
    subtitle_run = subtitle.add_run("v0.2 ｜ 对齐 V3 方案书 ｜ 2026 年 7 月 24 日 ｜ W2 专业底稿")
    set_run_font(subtitle_run, size=9, color=MID_GRAY)

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [24.2])
    write_cell(
        callout.cell(0, 0),
        "使用边界：本表用于把销售与收款循环的公开线索、资料缺口和计划阶段动作对应起来。它只提示“需要进一步了解什么”，不对违规、错报或交易真实性作定性，也不形成审计结论。R1、R2为当前开发资料上的工程版；R3—R8仍为候选规则，未接入时不得据此生成风险卡。",
        fill=LIGHT_BLUE,
        size=9,
    )

    add_heading(doc, "一、规则状态与使用口径")
    status_table = doc.add_table(rows=1, cols=3)
    status_table.style = "Table Grid"
    set_table_geometry(status_table, [3.0, 6.1, 15.1])
    headers = ["类别", "本表中的规则", "使用口径"]
    for index, label in enumerate(headers):
        write_cell(status_table.rows[0].cells[index], label, header=True, fill=BLUE, size=8.5)
    set_repeat_table_header(status_table.rows[0])
    status_rows = [
        ("工程版", "R1：应收—收入背离；R2：收入—经营现金流背离", "R1/R2可以在当前 DEV_T0 开发资料上由程序计算，但仍待专业口径复核。R2遇上年现金流很小、为零或正负号翻转时，同比百分比不宜直接作为判断依据，应改看金额变动、现金流/收入与人工解释。", LIGHT_GREEN),
        ("候选规则", "R3 产销量；R4 客户集中度；R5 总额/净额；R8 期末交易与期后回款", "仅作为资料检索、业务理解和后续索取资料的方向。资料、阈值、口径或专业复核未完成前，不生成风险卡。", LIGHT_AMBER),
        ("本表范围外", "R6 毛利率—行业；R7 管理层表述—数字", "这两条属于行业/披露维度，不是单一销售与收款环节。本表不将其伪装成已覆盖的业务循环规则。", LIGHT_GRAY),
    ]
    for category, rules, use, fill in status_rows:
        cells = status_table.add_row().cells
        write_cell(cells[0], category, fill=fill, size=8.4)
        write_cell(cells[1], rules, size=8.4)
        write_cell(cells[2], use, size=8.4)
    set_table_geometry(status_table, [3.0, 6.1, 15.1])

    add_heading(doc, "二、销售与收款循环映射矩阵")
    add_body(doc, "说明：公开年报未披露的信息统一表述为“目前未在公开资料中找到/需核验”，不据此推定公司从未发生相关事项。计划阶段动作是拟纳入进一步程序或资料索取清单，不表示已经执行函证、截止测试等审计程序。", size=8.8, color=MID_GRAY)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    widths = [2.2, 4.0, 3.15, 4.1, 4.2, 6.55]
    set_table_geometry(table, widths)
    headings = ["业务环节", "公开可观察信号 / 资料状态", "关联规则与当前状态", "待进一步了解的问题 / 候选认定", "正常解释候选", "资料缺口与计划阶段动作"]
    for index, heading in enumerate(headings):
        write_cell(table.rows[0].cells[index], heading, header=True, fill=BLUE, size=8.0)
    set_repeat_table_header(table.rows[0])

    rows = [
        (
            "1. 客户与信用政策管理",
            "前五大客户集中度、客户结构或信用政策出现跨期变化；如公开年报未披露客户身份或信用期，标为资料缺口。",
            "R4 候选；可为 R1/R2 的正常解释背景，但不是当前 R1/R2 的计算条件。",
            "是否需要进一步了解客户结构、信用期与回款节奏的变化。候选关注：收入发生、关联方关系及披露。",
            "新市场拓展；客户结构调整；行业账期整体变化。",
            "缺口：前五大客户/新增客户信息、信用政策与授权摘要。\n动作：了解变化原因；把客户背景、关联关系与信用政策列入拟索取资料。",
        ),
        (
            "2. 合同签订与订单管理",
            "收入政策或交易类型发生变化；公开披露的业务模式与同行口径存在需解释差异。",
            "R5 候选；合同资料也可作为 R1 触发后的补充解释资料。",
            "合同条款是否支持控制权转移、验收条件与结算安排；候选关注：准确性、截止、分类和列报。",
            "业务模式真实变化；合同条款差异；准则判断存在合理专业差异。",
            "缺口：主要合同关键条款、验收条件、结算安排。\n动作：选择代表性合同作为拟核验样本，核对收入政策与合同条款的对应关系。",
        ),
        (
            "3. 发货、履约与验收",
            "第四季度或年末收入占比异常时，需先确认是否有可比的季度数据和公开业务解释。",
            "R8 候选；若产销资料可得，可与 R3 候选联动。当前未接入程序。",
            "是否需要进一步了解期末确认、发货、验收与控制权转移的时点；候选关注：截止、发生。",
            "行业季节性；年末集中交付；大型项目在年末验收。",
            "缺口：发货、签收、验收、期末前后交易摘要。\n动作：将期末前后交易和支持性凭证列入拟截止核验范围。",
        ),
        (
            "4. 收入确认与开票",
            "R1：本年/上年收入与应收账款四字段增速背离。\nR2：收入正增长而经营现金流未同步改善。现金流基数很小、为零或正负号翻转时，不直接用同比百分比判断。",
            "R1 当前工程版；R2 当前工程版，R2 可比性保护和阈值仍待专业/工程验收。",
            "是否需要进一步了解收入、应收与现金回收之间的结算差异。候选关注：收入发生、准确性、截止及应收账款计价。",
            "信用政策或结算周期变化；行业账期拉长；新增客户或季节性回款。",
            "缺口：按客户/合同的收入—应收—回款勾稽、账龄明细、经营性应收应付变动。\n动作：保留来源、期间和公式；将差异原因与资料索取需求纳入风险卡草稿。",
        ),
        (
            "5. 应收账款与收款管理",
            "在账龄或周转资料可得时，关注账龄结构、期后回款和坏账准备变化；目前缺少字段时不得估算或自动计算。",
            "R1 增强方向（候选），不是当前四字段工程版的已实现条件。",
            "是否需要进一步了解可收回性、账龄迁徙与坏账准备；候选关注：应收账款存在、计价和权利义务。",
            "大客户固定付款周期；信用期延长；款项争议仍在正常协商。",
            "缺口：账龄明细、期后回款、主要客户结算条款、坏账准备依据。\n动作：列入拟函证与期后回款核验范围，并由人工决定是否执行。",
        ),
        (
            "6. 退货、折让与销售返利",
            "退货政策、历史退货率、期末后退货或返利信息披露不足，或存在需要共同解释的年末交易线索。",
            "R8 候选方向；不把退货/折让直接写成 R1/R2 的已实现触发信号。",
            "是否需要进一步了解期末发货后的退货、折让与预计退货处理；候选关注：截止、准确性和相关负债计量。",
            "行业惯例；合同约定的正常退货或返利条款；季节性促销。",
            "缺口：历史退货率、返利政策、期后退货/折让摘要及准备计提依据。\n动作：了解政策和历史波动，必要时列入期末前后交易的拟核验项目。",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            write_cell(cells[index], value, size=7.7)
    set_table_geometry(table, widths)

    add_heading(doc, "三、使用与复核要求")
    review_table = doc.add_table(rows=1, cols=2)
    review_table.style = "Table Grid"
    set_table_geometry(review_table, [5.2, 19.0])
    for index, label in enumerate(["复核点", "要求"]):
        write_cell(review_table.rows[0].cells[index], label, header=True, fill=BLUE, size=8.5)
    set_repeat_table_header(review_table.rows[0])
    review_rows = [
        ("来源", "每一条公司特定线索进入网页或风险卡前，必须补齐 evidence_id、资料名、披露日、PDF/印刷页和定位；未找到原件时标“待核验”。"),
        ("口径", "R1/R2计算只使用当前工程版登记字段。新增周转天数、客户集中度、退货率等指标前，先补字段定义、来源和异常样例。"),
        ("表述", "使用“待进一步了解”“资料依据不足”“建议索取”；不得使用定性指控、法律判断或结论性措辞。"),
        ("人工确认", "是否保留线索、是否索取内部资料、是否实施函证/截止测试等程序，均由审计人员决定并留痕。"),
    ]
    for label, detail in review_rows:
        cells = review_table.add_row().cells
        write_cell(cells[0], label, fill=LIGHT_GRAY, size=8.4)
        write_cell(cells[1], detail, size=8.4)
    set_table_geometry(review_table, [5.2, 19.0])

    add_body(doc, "版本记录：v0.2（2026-07-24）将 v0.1 的段落式初稿改为六环节映射矩阵；补充规则实施状态、认定、资料边界与 R2 可比性提示；保留 v0.1 原件以便追溯。", size=8.2, color=MID_GRAY)
    doc.core_properties.title = "销售与收款循环映射表 v0.2"
    doc.core_properties.subject = "审迹智链第二周专业底稿"
    doc.core_properties.author = "审迹智链项目组"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
