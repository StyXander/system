from __future__ import annotations

import importlib.machinery
import importlib.util
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parent
SOURCE = BASE_DIR / "02_最终确定方案" / "05_审迹智链_详细项目计划书_V2.4_资料入口指标库与快速MVP补充版.md"
# Fallback if rename has not happened yet
SOURCE_FALLBACK = BASE_DIR / "02_最终确定方案" / "05_审迹智链_详细项目计划书_V2_审计师定位修订版.md"
OUTPUT = BASE_DIR / "02_最终确定方案" / "05_审迹智链_详细项目计划书_V2.4_资料入口指标库与快速MVP补充版.docx"


def load_build_docs():
    py_path = BASE_DIR / "build_docs.py"
    if py_path.exists():
        sys.path.insert(0, str(BASE_DIR))
        import build_docs as mod  # type: ignore

        return mod
    pyc = BASE_DIR / "__pycache__" / "build_docs.cpython-312.pyc"
    loader = importlib.machinery.SourcelessFileLoader("build_docs", str(pyc))
    spec = importlib.util.spec_from_loader("build_docs", loader)
    mod = importlib.util.module_from_spec(spec)
    loader.exec_module(mod)
    return mod


bd = load_build_docs()
BLACK = bd.BLACK
DARK_BLUE = bd.DARK_BLUE
LIGHT_BLUE = bd.LIGHT_BLUE
MUTED = bd.MUTED
add_body = bd.add_body
add_callout = bd.add_callout
add_heading = bd.add_heading
add_hyperlink = bd.add_hyperlink
add_kicker = bd.add_kicker
add_list_item = bd.add_list_item
add_subtitle = bd.add_subtitle
add_table = bd.add_table
add_title = bd.add_title
add_two_column_fields = bd.add_two_column_fields
create_numbering = bd.create_numbering
set_doc_properties = bd.set_doc_properties
set_header_footer = bd.set_header_footer
set_run_font = bd.set_run_font
setup_document = bd.setup_document


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def table_widths(headers: list[str]) -> list[int]:
    count = len(headers)
    if count == 2:
        return [2350, 7010]
    if count == 3:
        if headers[0] in {"组别", "时间"}:
            return [1300, 3100, 4960]
        if headers[0] == "类别":
            return [1650, 3310, 4400]
        return [1900, 3260, 4200]
    if count == 4:
        if headers[0] == "编号":
            return [700, 3800, 2310, 2550]
        return [1500, 2720, 2570, 2570]
    if count == 5:
        return [2600, 1690, 1690, 1690, 1690]
    base = 9360 // count
    widths = [base] * count
    widths[-1] += 9360 - sum(widths)
    return widths


def add_inline(paragraph, text: str, *, size: float | None = None, color: str = BLACK):
    pattern = re.compile(r"(\*\*.+?\*\*|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            set_run_font(paragraph.add_run(text[pos : match.start()]), size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), size=size, bold=True, color=DARK_BLUE)
        else:
            link = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", token)
            assert link is not None
            label, target = link.groups()
            if target.startswith(("http://", "https://")):
                add_hyperlink(paragraph, label, target)
            else:
                set_run_font(paragraph.add_run(label), size=size, color=DARK_BLUE)
        pos = match.end()
    if pos < len(text):
        set_run_font(paragraph.add_run(text[pos:]), size=size, color=color)


def add_rich_body(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal")
    add_inline(p, text)
    return p


def add_rich_list_item(doc: Document, text: str, num_id: int, *, line_spacing: float):
    p = add_list_item(doc, "", num_id, line_spacing=line_spacing, after=4)
    add_inline(p, text)
    return p


def add_project_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = add_table(
        doc,
        headers,
        rows,
        table_widths(headers),
        header_fill=LIGHT_BLUE,
        font_size=8.6 if len(headers) >= 4 else 8.9,
        cell_vmargin=70,
    )
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=8.6 if len(headers) >= 4 else 8.9, bold=True, color=DARK_BLUE)
    return table


def parse_markdown_body(doc: Document, lines: list[str]):
    bullet_id = create_numbering(doc, kind="bullet", text_indent=540, hanging=271)
    current_number_id: int | None = None
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if not stripped:
            current_number_id = None
            i += 1
            continue

        if stripped.startswith("# "):
            i += 1
            continue

        if stripped.startswith(">"):
            quote_parts: list[str] = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_parts.append(lines[i].strip().lstrip(">").strip())
                i += 1
            text = " ".join(quote_parts).replace("**", "")
            add_callout(doc, "核心表述：", text, fill="F4F6F9", accent=DARK_BLUE)
            current_number_id = None
            continue

        if stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), 1)
            current_number_id = None
            i += 1
            continue

        if stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), 2)
            current_number_id = None
            i += 1
            continue

        if stripped.startswith("#### "):
            add_heading(doc, stripped[5:].strip(), 3)
            current_number_id = None
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1].strip()):
            headers = split_table_row(stripped)
            rows: list[list[str]] = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = split_table_row(lines[i].strip())
                if len(row) == len(headers):
                    rows.append([cell.replace("**", "") for cell in row])
                i += 1
            add_project_table(doc, [h.replace("**", "") for h in headers], rows)
            current_number_id = None
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            add_rich_list_item(doc, bullet.group(1), bullet_id, line_spacing=1.25)
            current_number_id = None
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            if current_number_id is None:
                current_number_id = create_numbering(doc, kind="decimal", text_indent=540, hanging=271)
            add_rich_list_item(doc, numbered.group(1), current_number_id, line_spacing=1.25)
            i += 1
            continue

        add_rich_body(doc, stripped)
        current_number_id = None
        i += 1


def parse_meta(lines: list[str]) -> dict[str, str]:
    meta: dict[str, str] = {}
    for line in lines:
        m = re.match(r">\s*\*\*(.+?)：\*\*(.+)$", line.strip())
        if m:
            meta[m.group(1).strip()] = m.group(2).strip()
    return meta


def build() -> Path:
    source = SOURCE if SOURCE.exists() else SOURCE_FALLBACK
    if not source.exists():
        raise FileNotFoundError(f"source markdown not found: {SOURCE} or {SOURCE_FALLBACK}")

    lines = source.read_text(encoding="utf-8").splitlines()
    meta = parse_meta(lines)
    first_section = next((i for i, line in enumerate(lines) if line.startswith("## ")), 0)

    version = meta.get("版本", "V2.4.2·可运行主链、智能体与RAG必做补充版")
    date = meta.get("日期", "2026年7月21日")
    status = meta.get(
        "当前状态",
        "指导老师已同意团队开始推进项目；该决定表示项目可以进入执行准备与最小原型阶段，不等同于8条规则、案例、T0/T1数据包、实验配置或效果结论已经确认或完成。",
    )

    doc = Document()
    section = setup_document(doc, "compact_reference_guide")
    set_header_footer(section, "审迹智链 | 审计前置阶段", "详细项目计划书 V2.4.2")
    set_doc_properties(
        doc,
        title="审迹智链详细项目计划书（V2.4.2可运行主链、智能体与RAG必做补充版）",
        subject="面向审计业务承接、续聘和计划阶段的收入确认重大错报风险预筛查方案",
    )
    doc.core_properties.keywords = "数智会计, 审计计划, 业务承接, 收入确认, 交叉验证, 反证挑战, RAG, AuditTrace, V2.4.2"

    add_kicker(doc, "2026 北京市大学生数智会计创新应用竞赛", align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    add_title(doc, "审迹智链（AuditTrace）", size=29, align=WD_ALIGN_PARAGRAPH.CENTER, after=7)
    add_subtitle(
        doc,
        "面向审计前置阶段的\n多维交叉验证与证据缺口识别智能体",
        size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=6,
    )
    add_subtitle(
        doc,
        "基于公开资料、证据约束与反证挑战的收入确认重大错报风险预筛查",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=22,
    )
    add_two_column_fields(
        doc,
        [
            ("核心用户：", "会计师事务所审计项目组"),
            ("核心阶段：", "业务承接 / 续聘 / 审计计划"),
            ("版本：", version),
            ("日期：", date),
        ],
    )
    add_callout(
        doc,
        "当前阶段：",
        status,
        fill="F4F6F9",
        accent=DARK_BLUE,
    )
    cover_note = doc.add_paragraph()
    cover_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_note.paragraph_format.space_before = Pt(22)
    set_run_font(
        cover_note.add_run("团队执行底稿 · V2.4.2 可运行主链 / 智能体 / RAG必做补充 · 尚无正式实验结果"),
        size=9.5,
        color=MUTED,
    )
    cover_note.add_run().add_break(WD_BREAK.PAGE)

    parse_markdown_body(doc, lines[first_section:])
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
