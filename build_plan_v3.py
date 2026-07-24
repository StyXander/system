"""将审迹智链 V3 方案书 Markdown 生成同名 Word 文档。

文档采用 narrative_proposal 版式：Letter 纵向、1 英寸页边距、
Microsoft YaHei 正文、蓝色层级标题、固定宽度表格和页脚页码。
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "02_最终确定方案" / "05_审迹智链_项目方案书_V3_全面复盘修订版.md"
OUTPUT = ROOT / "02_最终确定方案" / "05_审迹智链_项目方案书_V3_全面复盘修订版.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = RGBColor(88, 98, 112)
CONTENT_WIDTH = 9360


def set_run_font(run, size: float, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def table_widths(column_count: int) -> list[int]:
    patterns = {
        2: [2400, 6960],
        3: [2520, 3330, 3510],
        4: [1900, 3000, 2000, 2460],
        5: [1500, 2200, 1700, 2460, 1500],
    }
    return patterns.get(column_count, [CONTENT_WIDTH // column_count] * column_count)


def clean_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    text = text.replace("`", "")
    text = text.replace("**", "")
    return text.strip()


def add_text(paragraph, text: str, size: float = 10.5, color: str | None = None) -> None:
    """支持粗体片段，链接在Word中保留为普通来源文字。"""
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > cursor:
            run = paragraph.add_run(clean_inline(text[cursor:match.start()]))
            set_run_font(run, size, color=color)
        run = paragraph.add_run(clean_inline(match.group(1)))
        set_run_font(run, size, bold=True, color=color)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(clean_inline(text[cursor:]))
        set_run_font(run, size, color=color)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_paragraph(paragraph, after: float = 6, before: float = 0, line: float = 1.25) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(after)
    fmt.space_before = Pt(before)
    fmt.line_spacing = line
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def add_table(doc: Document, rows: list[list[str]]) -> None:
    column_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    set_table_widths(table, table_widths(column_count))
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.text = ""
            paragraph = cell.paragraphs[0]
            set_paragraph(paragraph, after=1, line=1.1)
            add_text(paragraph, value, size=8.2 if column_count >= 4 else 8.7, color=NAVY if row_index == 0 else None)
            if row_index == 0:
                shade_cell(cell, LIGHT_BLUE)
                for run in paragraph.runs:
                    run.bold = True
            elif row_index % 2 == 0:
                shade_cell(cell, LIGHT_GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (15.5, NAVY, 16, 8),
        "Heading 2": (12.5, BLUE, 12, 6),
        "Heading 3": (11.5, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.text = "审迹智链（AuditTrace）｜V3 项目方案书"
    set_paragraph(header, after=0, line=1.0)
    for run in header.runs:
        set_run_font(run, 8.5, color="667085")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "北京市大学生数智会计创新应用竞赛｜第 "
    add_page_field(footer)
    footer.add_run(" 页")
    set_paragraph(footer, after=0, line=1.0)
    for run in footer.runs:
        set_run_font(run, 8.5, color="667085")


def add_cover(doc: Document, title: str, metadata: list[str]) -> None:
    for _ in range(8):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(kicker, "北京市大学生数智会计创新应用竞赛", size=11, color=BLUE)
    set_paragraph(kicker, after=14, line=1.0)
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(heading, title, size=24, color=NAVY)
    for run in heading.runs:
        run.bold = True
    set_paragraph(heading, after=12, line=1.15)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(subtitle, "面向审计前置阶段的证据约束型收入风险预筛查多智能体", size=13, color="475467")
    set_paragraph(subtitle, after=34, line=1.2)
    for item in metadata:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(paragraph, item, size=10.5, color="475467")
        set_paragraph(paragraph, after=5, line=1.1)
    doc.add_page_break()


def build() -> None:
    markdown = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    title = markdown[0].lstrip("# ").strip()
    metadata = [line[1:].strip() for line in markdown[1:] if line.startswith(">")]
    add_cover(doc, title, metadata)

    index = 0
    in_code = False
    code_lines: list[str] = []
    while index < len(markdown):
        line = markdown[index]
        stripped = line.strip()
        if index == 0 or line.startswith(">"):
            index += 1
            continue
        if stripped.startswith("```"):
            if in_code:
                table = doc.add_table(rows=1, cols=1)
                set_table_widths(table, [CONTENT_WIDTH])
                cell = table.cell(0, 0)
                shade_cell(cell, "F4F6F9")
                cell.text = ""
                paragraph = cell.paragraphs[0]
                set_paragraph(paragraph, after=1, line=1.08)
                add_text(paragraph, "\n".join(code_lines), size=8.7, color=NAVY)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if line.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 3")
            add_text(paragraph, line[4:], size=11.5, color=NAVY)
            index += 1
            continue
        if line.startswith("## "):
            paragraph = doc.add_paragraph(style="Heading 1")
            add_text(paragraph, line[3:], size=15.5, color=NAVY)
            index += 1
            continue
        if line.startswith("# "):
            index += 1
            continue
        if line.startswith("|") and "|" in line[1:]:
            rows: list[list[str]] = []
            while index < len(markdown) and markdown[index].startswith("|"):
                candidate = markdown[index].strip()
                if not re.fullmatch(r"\|[\s|:-]+\|", candidate):
                    rows.append([cell.strip() for cell in candidate.strip("|").split("|")])
                index += 1
            if rows:
                add_table(doc, rows)
            continue
        match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if match:
            paragraph = doc.add_paragraph(style="List Number")
            set_paragraph(paragraph, after=4, line=1.2)
            add_text(paragraph, match.group(2), size=10.5)
            index += 1
            continue
        paragraph = doc.add_paragraph()
        set_paragraph(paragraph, after=6, line=1.25)
        add_text(paragraph, stripped, size=10.5)
        index += 1

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
